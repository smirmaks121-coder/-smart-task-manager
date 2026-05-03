from fastapi import FastAPI, Request, Depends, Form, Cookie
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from datetime import datetime
from docx import Document
from passlib.context import CryptContext
from jose import jwt
import models
from database import SessionLocal, engine

# Создаем таблицы в БД
models.Base.metadata.create_all(bind=engine)

app = FastAPI()

# Подключаем статику и шаблоны
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# Настройки безопасности
SECRET_KEY = "super-secret-key-for-manager"
ALGORITHM = "HS256"
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Зависимость для БД
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Получение текущего юзера
async def get_current_user(request: Request, db: Session = Depends(get_db)):
    token = request.cookies.get("access_token")
    if not token:
        return None
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        return db.query(models.User).filter(models.User.username == username).first()
    except:
        return None

# --- МАРШРУТЫ АВТОРИЗАЦИИ ---

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    return templates.TemplateResponse(
        request=request,
        name="login.html",
        context={"request": request}
    )

@app.post("/register")
async def register(username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    hashed = pwd_context.hash(password)
    new_user = models.User(username=username, hashed_password=hashed)
    db.add(new_user)
    db.commit()
    return RedirectResponse(url="/login", status_code=303)

@app.post("/login")
async def login(username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.username == username).first()
    if not user or not pwd_context.verify(password, user.hashed_password):
        return RedirectResponse(url="/login?error=1", status_code=303)
    
    token = jwt.encode({"sub": user.username}, SECRET_KEY, algorithm=ALGORITHM)
    response = RedirectResponse(url="/", status_code=303)
    response.set_cookie(key="access_token", value=token, httponly=True)
    return response

@app.get("/logout")
async def logout():
    response = RedirectResponse(url="/login")
    response.delete_cookie("access_token")
    return response

# --- ОСНОВНАЯ ЛОГИКА ЗАДАЧ ---

@app.get("/", response_class=HTMLResponse)
async def home(request: Request, db: Session = Depends(get_db), user: models.User = Depends(get_current_user)):
    if not user:
        return RedirectResponse(url="/login")
    
    tasks = db.query(models.Task).filter(models.Task.owner_id == user.id, models.Task.is_deleted == False).order_by(models.Task.due_date).all()
    trash = db.query(models.Task).filter(models.Task.owner_id == user.id, models.Task.is_deleted == True).all()
    
    total = len(tasks)
    completed = len([t for t in tasks if t.is_completed])
    progress = int((completed / total) * 100) if total > 0 else 0
    
    return templates.TemplateResponse(
        request=request,
        name="index.html",
        context={"request": request, "tasks": tasks, "trash": trash, "progress": progress, "user": user}
    )

@app.post("/add")
async def add_task(title: str = Form(...), category: str = Form(...), due_date: str = Form(...), db: Session = Depends(get_db), user: models.User = Depends(get_current_user)):
    if not user: return RedirectResponse(url="/login")
    date_obj = datetime.strptime(due_date, '%Y-%m-%d').date()
    db.add(models.Task(title=title, category=category, due_date=date_obj, owner_id=user.id))
    db.commit()
    return RedirectResponse(url="/", status_code=303)

@app.post("/toggle/{task_id}")
async def toggle_task(task_id: int, db: Session = Depends(get_db), user: models.User = Depends(get_current_user)):
    task = db.query(models.Task).filter(models.Task.id == task_id, models.Task.owner_id == user.id).first()
    if task:
        task.is_completed = not task.is_completed
        db.commit()
    return RedirectResponse(url="/", status_code=303)

@app.post("/delete/{task_id}")
async def to_trash(task_id: int, db: Session = Depends(get_db), user: models.User = Depends(get_current_user)):
    task = db.query(models.Task).filter(models.Task.id == task_id, models.Task.owner_id == user.id).first()
    if task:
        task.is_deleted = True
        db.commit()
    return RedirectResponse(url="/", status_code=303)

@app.get("/export/word")
async def export_to_word(db: Session = Depends(get_db), user: models.User = Depends(get_current_user)):
    if not user: return RedirectResponse(url="/login")
    tasks = db.query(models.Task).filter(models.Task.owner_id == user.id, models.Task.is_deleted == False).all()
    doc = Document()
    doc.add_heading(f'Задачи: {user.username}', 0)
    for t in tasks:
        status = "[V]" if t.is_completed else "[ ]"
        doc.add_paragraph(f"{t.due_date} | {t.title} ({t.category}) {status}")
    
    file_path = f"tasks_{user.username}.docx"
    doc.save(file_path)
    return FileResponse(file_path, filename="My_Tasks.docx")
